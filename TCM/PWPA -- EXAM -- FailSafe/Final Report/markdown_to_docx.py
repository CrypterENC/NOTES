#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    """Add a horizontal line to a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def parse_markdown_to_docx(md_file, docx_file):
    """Convert Markdown report to Word document"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            heading_text = line[2:].strip()
            p = doc.add_heading(heading_text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            p = doc.add_heading(heading_text, level=2)
            i += 1
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_heading(heading_text, level=3)
            i += 1
        elif line.startswith('#### '):
            heading_text = line[5:].strip()
            p = doc.add_heading(heading_text, level=4)
            i += 1
        
        # Handle horizontal rules
        elif line.strip() == '---':
            p = doc.add_paragraph()
            add_horizontal_line(p)
            i += 1
        
        # Handle code blocks
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='List Bullet')
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            
            # Format as code
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            i += 1  # Skip closing ```
        
        # Handle tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Parse table
            header_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            i += 1
            
            # Skip separator row
            if '---' in lines[i]:
                i += 1
            
            # Create table
            table = doc.add_table(rows=1, cols=len(header_cells))
            table.style = 'Light Grid Accent 1'
            
            # Add header
            header_cells_obj = table.rows[0].cells
            for idx, cell_text in enumerate(header_cells):
                header_cells_obj[idx].text = cell_text.replace('**', '').replace('*', '')
            
            # Add rows
            while i < len(lines) and '|' in lines[i]:
                row_cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                if len(row_cells) == len(header_cells):
                    row = table.add_row()
                    for idx, cell_text in enumerate(row_cells):
                        row.cells[idx].text = cell_text.replace('**', '').replace('*', '').replace('`', '')
                i += 1
        
        # Handle bold and italic text
        elif '**' in line or '*' in line or '`' in line:
            p = doc.add_paragraph()
            
            # Parse inline formatting
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                elif part:
                    p.add_run(part)
            
            i += 1
        
        # Handle bullet points
        elif line.strip().startswith('- '):
            bullet_text = line.strip()[2:]
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1
        
        # Handle numbered lists
        elif line.strip() and line.strip()[0].isdigit() and '.' in line[:3]:
            list_text = line.strip()
            match = re.match(r'^(\d+)\.\s+(.*)', list_text)
            if match:
                p = doc.add_paragraph(match.group(2), style='List Number')
                i += 1
            else:
                p = doc.add_paragraph(line)
                i += 1
        
        # Regular paragraphs
        else:
            if line.strip():
                p = doc.add_paragraph(line)
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Word document created successfully!")
    print(f"Location: {docx_file}")

if __name__ == '__main__':
    md_file = r'c:\Users\mathe\OneDrive\Documents\ObsidianVault\notes\NOTES\TCM\PWPA -- EXAM -- FailSafe\Final Report\FailSafe - Final Penetration Test Report.md'
    docx_file = r'c:\Users\mathe\OneDrive\Documents\ObsidianVault\notes\NOTES\TCM\PWPA -- EXAM -- FailSafe\Final Report\FailSafe - Final Penetration Test Report.docx'
    
    parse_markdown_to_docx(md_file, docx_file)
